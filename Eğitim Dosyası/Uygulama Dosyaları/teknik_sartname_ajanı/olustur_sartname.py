# -*- coding: utf-8 -*-
"""
Personel Servisi Hizmet Alimi Teknik Sartnamesi
Word Belgesi Olusturucu
Kalite Kontrolcu Ajani tarafindan uretilmistir.
Tarih: 24.04.2026
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ─────────────────────────────────────────
# YARDIMCI FONKSİYONLAR
# ─────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """Hücre arka plan rengini ayarla."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(table):
    """Tabloya ince çerçeve ekle."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '999999')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def add_header_row(table, headers, bg_color='1F5C99'):
    """Tablo başlık satırı ekle."""
    row = table.rows[0]
    for i, header in enumerate(headers):
        cell = row.cells[i]
        cell.text = header
        set_cell_bg(cell, bg_color)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0] if para.runs else para.add_run(header)
        if not para.runs:
            cell.text = ''
            run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)

def add_table_row(table, row_index, values, bold_first=False, alt_row=False):
    """Normal tablo satırı ekle."""
    row = table.rows[row_index]
    bg = 'EBF3FB' if alt_row else 'FFFFFF'
    for i, val in enumerate(values):
        cell = row.cells[i]
        set_cell_bg(cell, bg)
        cell.text = str(val)
        para = cell.paragraphs[0]
        if i == 0 and bold_first:
            for run in para.runs:
                run.bold = True
        for run in para.runs:
            run.font.size = Pt(9)

def add_heading(doc, text, level, color=None):
    """Başlık ekle."""
    para = doc.add_heading(text, level=level)
    if color:
        for run in para.runs:
            run.font.color.rgb = color
    return para

def add_bullet(doc, text, level=0):
    """Madde işaretli paragraf ekle."""
    para = doc.add_paragraph(style='List Bullet')
    run = para.add_run(text)
    run.font.size = Pt(10)
    return para

def add_numbered(doc, text):
    """Numaralı liste ekle."""
    para = doc.add_paragraph(style='List Number')
    run = para.add_run(text)
    run.font.size = Pt(10)
    return para

def add_body(doc, text):
    """Normal paragraf ekle."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(10)
    return para

def add_note(doc, text):
    """Not/uyarı paragrafı ekle (italik, gri)."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return para

def add_section_break_line(doc):
    """Bölüm ayraç çizgisi."""
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F5C99')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para

# ─────────────────────────────────────────
# BELGE OLUŞTURMA
# ─────────────────────────────────────────

doc = Document()

# Sayfa ayarları
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(3)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)

# Varsayılan font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

# ────────────────────────────
# KAPAK SAYFASI
# ────────────────────────────
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_para.add_run('PERSONEL SERVİSİ HİZMET ALIMI')
title_run.bold = True
title_run.font.size = Pt(20)
title_run.font.color.rgb = RGBColor(0x1F, 0x5C, 0x99)
title_run.font.name = 'Calibri'

subtitle_para = doc.add_paragraph()
subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle_para.add_run('TEKNİK ŞARTNAME')
subtitle_run.bold = True
subtitle_run.font.size = Pt(16)
subtitle_run.font.color.rgb = RGBColor(0x1F, 0x5C, 0x99)
subtitle_run.font.name = 'Calibri'

doc.add_paragraph()

meta_table = doc.add_table(rows=5, cols=2)
meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_data = [
    ('Hazırlık Tarihi', '24.04.2026'),
    ('Versiyon', '1.1'),
    ('Durum', 'Kalite Kontrol Onaylı'),
    ('Hazırlayan Birim', 'Teknik Şartname Uzmanı + Hukuk Müşaviri'),
    ('Kalite Kontrol', 'Kalite Kontrolcü Ajanı'),
]
for i, (k, v) in enumerate(meta_data):
    row = meta_table.rows[i]
    row.cells[0].text = k
    row.cells[1].text = v
    set_cell_bg(row.cells[0], 'D6E4F0')
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)
                run.font.name = 'Calibri'
    row.cells[0].paragraphs[0].runs[0].bold = True
set_cell_borders(meta_table)

doc.add_page_break()

# ────────────────────────────
# BÖLÜM 0 — TANIMLAR
# ────────────────────────────
add_section_break_line(doc)
h = doc.add_heading('BÖLÜM 0 — TANIMLAR', level=1)
h.runs[0].font.color.rgb = RGBColor(0x1F, 0x5C, 0x99)

add_body(doc, 'Bu şartnamede geçen terimler aşağıdaki anlamlarda kullanılır:')
doc.add_paragraph()

terms = [
    ('İdare', 'Hizmet alımını gerçekleştiren kamu kurum veya kuruluşu'),
    ('Yüklenici', 'Personel servisi hizmetini sunan ve sözleşme imzalayan gerçek veya tüzel kişi'),
    ('Sefer', 'Belirli bir güzergah üzerinde planlanan tek yönlü araç seyahati'),
    ('Güzergah', 'Başlangıç noktasından bitiş noktasına kadar onaylanmış durakları kapsayan rota'),
    ('Aktif Filo', 'Günlük seferlerde görev yapan araçlar topluluğu'),
    ('Yedek Araç', 'Arıza veya acil durumlarda aktif filonun yerine geçmek üzere hazır tutulan araç'),
    ('Durak', 'Yolcuların bineceği veya ineceği, sözleşme ekinde tanımlanmış güzergah noktası'),
    ('Merkez Operatörü', 'Yüklenici bünyesinde seferleri gerçek zamanlı koordine eden personel'),
    ('KPI', 'Anahtar Performans Göstergesi; hizmet kalitesini ölçen sayısal hedef'),
    ('Telematik Sistem', 'GPS, sürücü davranış sensörü ve merkez iletişimini birleştiren araç izleme altyapısı'),
    ('İhlal Tespit Tutanağı', 'İdarece düzenlenen, yüklenicinin yükümlülüğünü yerine getirmediğini belgeleyen resmi tutanak'),
]
tbl = doc.add_table(rows=len(terms)+1, cols=2)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
add_header_row(tbl, ['Terim', 'Tanım'])
for i, (term, defn) in enumerate(terms):
    add_table_row(tbl, i+1, [term, defn], bold_first=True, alt_row=(i%2==0))
set_cell_borders(tbl)

# ────────────────────────────
# BÖLÜM 1 — ARAÇ TEKNİK ÖZELLİKLERİ
# ────────────────────────────
doc.add_page_break()
add_section_break_line(doc)
h = doc.add_heading('BÖLÜM 1 — ARAÇ TEKNİK ÖZELLİKLERİ', level=1)
h.runs[0].font.color.rgb = RGBColor(0x1F, 0x5C, 0x99)

# 1.1
doc.add_heading('1.1 Araç Tipi ve Kapasitesi', level=2)
add_body(doc, 'Yüklenici, güzergah ihtiyaçlarına göre aşağıda tanımlanan araç tiplerinden oluşan bir filo temin etmekle yükümlüdür.')
doc.add_paragraph()

arac_tipleri = [
    ('Minibüs', '14 + 1 = 15 kişi', 'Kısa mesafe ve dar güzergahlar'),
    ('Midibüs', '25 + 1 = 26 kişi', 'Ana hat ve yüksek kapasiteli güzergahlar'),
]
t = doc.add_table(rows=3, cols=3)
t.style = 'Table Grid'
add_header_row(t, ['Araç Tipi', 'Yolcu Kapasitesi (Yolcu + Sürücü)', 'Kullanım Amacı'])
for i, row_data in enumerate(arac_tipleri):
    add_table_row(t, i+1, row_data, alt_row=(i%2==0))
set_cell_borders(t)
doc.add_paragraph()

doc.add_heading('Araç Yaşı ve Model Yılı Şartları:', level=3)
add_bullet(doc, 'Araçların tescil tarihinden itibaren yaşı en fazla 5 (beş) yıl olacaktır.')
add_bullet(doc, 'Tüm araçlar en az 2020 model ve üzeri olacaktır.')
add_bullet(doc, 'Sözleşme süresi boyunca araç yaşı 5 yılı aştığında yüklenici, idarenin yazılı talebi olmaksızın araç yenileme yükümlülüğünü 30 gün içinde yerine getirecektir.')
add_bullet(doc, 'Araçların tescil belgeleri, muayene belgeleri ve sigorta poliçeleri yüklenici tarafından sözleşme imzalanmadan önce idareye sunulacaktır.')

doc.add_heading('Yakıt Tipi ve Emisyon Standardı:', level=3)
add_bullet(doc, 'Araçlar dizel, LPG veya hibrit yakıt sistemine sahip olabilir.')
add_bullet(doc, 'LPG dönüşümü yapılmış araçların Karayolları Trafik Yönetmeliği kapsamında dönüşüm belgesi ve LPG araç muayenesi bulunmalıdır.')
add_bullet(doc, 'Tüm araçlar minimum Euro 5 emisyon standardını karşılamalıdır; Euro 5 altındaki araçlar sefere çıkarılamaz.')

doc.add_heading('Araç Sayısı:', level=3)
add_bullet(doc, 'Aktif sefer filosunda çalışacak araç sayısı idare tarafından belirlenecek olup sözleşme ekinde listelenecektir.')
add_bullet(doc, 'Aktif filonun %15\'i oranında yedek araç her an hazır bulundurulacaktır (minimum 1 adet).')

# 1.2
doc.add_heading('1.2 Zorunlu Güvenlik Donanımı', level=2)

doc.add_heading('Aktif Güvenlik Sistemleri:', level=3)
add_bullet(doc, 'ABS (Kilitlenme Önleyici Fren Sistemi)')
add_bullet(doc, 'ESP (Elektronik Kararlılık Programı)')
add_bullet(doc, 'EBD (Elektronik Fren Güç Dağıtımı)')
add_bullet(doc, 'Çarpışma uyarı sistemi (opsiyonel artı puan olarak değerlendirilir)')

doc.add_heading('Pasif Güvenlik Ekipmanları:', level=3)
add_bullet(doc, 'Yolcu emniyet kemeri — her koltukta, çalışır durumda, TSE belgeli')
add_bullet(doc, 'Yangın söndürücü — 2 kg kuru kimyevi toz, TS EN 3 belgeli, son kullanma tarihi geçmemiş')
add_bullet(doc, 'İlk yardım çantası — TS 11888 standardına uygun, içeriği tam ve geçerli tarihli')
add_bullet(doc, 'Reflektif güvenlik yeleği — en az 2 adet, her an erişilebilir konumda')
add_bullet(doc, 'Uyarı üçgeni — 2 adet, TS 3127 normuna uygun')
add_bullet(doc, 'Çekici halat veya çekme demiri — araç tipine uygun kapasitede')

doc.add_heading('Görüş ve Manevra Güvenliği:', level=3)
add_bullet(doc, 'Arka görüş kamerası — geri vites anında aktif, sürücü ekranına entegre')
add_bullet(doc, 'Geri vites sesli uyarı sistemi — dış uyarı sesi min. 85 dB')

doc.add_heading('Sigorta Yükümlülükleri:', level=3)
add_bullet(doc, 'Zorunlu Mali Sorumluluk (Trafik) Sigortası: Her araç için yasal asgari teminat tutarlarını karşılayan geçerli poliçe zorunludur.')
add_bullet(doc, 'Kasko Sigortası: Aracın piyasa değerini karşılayacak kapsamda; her araç için ayrı poliçe.')
add_bullet(doc, 'Ferdi Kaza Sigortası: Her koltuk başına kaza teminatı; sürücü dahil tüm yolcuları kapsar.')
add_bullet(doc, 'Tüm sigorta poliçelerinin birer sureti sözleşme imzalanmadan önce ve her yenileme sonrasında 5 iş günü içinde idareye teslim edilecektir.')
add_bullet(doc, 'Poliçe süresinin dolmasından en az 15 gün önce yüklenici yenileme işlemini tamamlayacaktır.')

# 1.3
doc.add_heading('1.3 Konfor ve Yolcu Donanımı', level=2)

doc.add_heading('İklimlendirme:', level=3)
add_bullet(doc, 'Klima sistemi: ön ve arka bölge bağımsız kontrollü; dış sıcaklık +35°C\'de iç sıcaklığı +22°C\'ye düşürebilmeli')
add_bullet(doc, 'Isıtma sistemi: dış sıcaklık -10°C\'de iç sıcaklık minimum +18°C sağlamalı')

doc.add_heading('Yolcu Konforu:', level=3)
add_bullet(doc, 'USB şarj noktaları: her koltuk başına en az 1 adet (Type-A veya Type-C)')
add_bullet(doc, 'Perdeler veya güneşlik filmleri: cam kararma oranı %30–60 arasında, TS EN 1 standartlarına uygun')
add_bullet(doc, 'Koltuk kılıfları: kumaş veya deri kaplama, haftalık değişim ve yıkama zorunluluğu')
add_bullet(doc, 'Zemin kaplaması: anti-slip (kaymaz) özellikte, kolay temizlenebilir malzeme (kauçuk veya PVC esaslı)')
add_bullet(doc, 'Koltuk düzeni: ergonomik, sırt dayanaklı; midibüs araçlarında bagaj rafı bulunacaktır')

doc.add_heading('Erişilebilirlik:', level=3)
add_bullet(doc, 'Araçlarda en az 1 adet el tutamağı veya basamak yardımcısı bulunacaktır.')
add_bullet(doc, 'Engelli bireylere yönelik zorunlu düzenlemeler ilgili mevzuat çerçevesinde uygulanacaktır.')

# 1.4
doc.add_heading('1.4 İletişim ve Takip Sistemleri', level=2)

doc.add_heading('GPS ve Telematik Takip:', level=3)
add_bullet(doc, 'Her araçta 4.5G (LTE) destekli GPS takip cihazı zorunludur.')
add_bullet(doc, 'Gerçek zamanlı konum bildirimi: 30 saniye güncelleme aralığı.')
add_bullet(doc, 'GPS verisi kesintisiz olarak yüklenicinin merkez sistemine ve idare izleme portalına iletilecektir.')
add_bullet(doc, 'Araç kapalıyken son konum bilgisi sistemde 72 saat saklanacaktır.')

doc.add_heading('Sürücü Davranış İzleme:', level=3)
add_bullet(doc, 'Ani fren (0,4G üzeri yavaşlama)')
add_bullet(doc, 'Sert dönüş (0,3G üzeri yatay ivme)')
add_bullet(doc, 'Hız aşımı (güzergah hız limitini aşma)')
add_bullet(doc, 'Aşırı hızlanma (0–50 km/h arası 5 saniyenin altı)')
add_note(doc, 'Aylık sürücü davranış raporu yüklenici tarafından idareye sunulacaktır.')

doc.add_heading('Araç İçi Kamera Sistemi:', level=3)
add_bullet(doc, 'En az 2 kamera açısı: sürücü kabini + yolcu bölümü')
add_bullet(doc, 'Kayıt süresi: 30 gün sürekli kayıt, üzerine yazma korumalı')
add_bullet(doc, 'Kamera görüntüleri yalnızca yetkilendirilmiş idare personeli tarafından talep üzerine incelenebilir.')
add_bullet(doc, 'KVKK kapsamındaki aydınlatma yükümlülükleri yüklenici tarafından yerine getirilecektir.')

# ────────────────────────────
# BÖLÜM 2 — SÜRÜCÜ YETERLİLİK STANDARTLARI
# ────────────────────────────
doc.add_page_break()
add_section_break_line(doc)
h = doc.add_heading('BÖLÜM 2 — SÜRÜCÜ YETERLİLİK STANDARTLARI', level=1)
h.runs[0].font.color.rgb = RGBColor(0x1F, 0x5C, 0x99)

doc.add_heading('2.1 Zorunlu Belgeler ve Şartlar', level=2)
add_body(doc, 'Yüklenici bünyesinde çalışacak tüm sürücülerin aşağıdaki belgelere sahip olması zorunludur.')
doc.add_paragraph()

surucuBelge = [
    ('D2 Sürücü Belgesi', 'Minibüs (14+1) kullananlar', 'Güncel ve geçerli'),
    ('D Sınıfı Sürücü Belgesi', 'Midibüs (25+1) kullananlar', 'Güncel ve geçerli'),
    ('SRC-3 Mesleki Yeterlilik Belgesi (Yolcu Taşımacılığı)', 'Tüm sürücüler', 'Güncel ve geçerli'),
    ('Psikoteknik Değerlendirme Raporu', 'Tüm sürücüler', 'Son 2 yıl içinde alınmış'),
    ('Adli Sicil Kaydı (Sabıka)', 'Tüm sürücüler', 'Son 3 ay içinde alınmış'),
    ('Sağlık Raporu', 'Tüm sürücüler (göz muayenesi dahil)', 'Son 1 yıl içinde alınmış'),
]
t = doc.add_table(rows=len(surucuBelge)+1, cols=3)
t.style = 'Table Grid'
add_header_row(t, ['Belge / Şart', 'Kapsam', 'Geçerlilik Süresi'])
for i, row_data in enumerate(surucuBelge):
    add_table_row(t, i+1, row_data, alt_row=(i%2==0))
set_cell_borders(t)
doc.add_paragraph()

doc.add_heading('Ek Şartlar:', level=3)
add_bullet(doc, 'Sürücüler 18 yaşından büyük, 65 yaşından küçük olacaktır.')
add_bullet(doc, 'Aktif alkol/uyuşturucu kullanımı bulunmamalıdır (periyodik denetim hakkı idareye aittir).')
add_bullet(doc, 'Son 5 yılda sürücü belgesine 6 veya daha fazla ceza puanı işlenmemiş olmalıdır.')

doc.add_heading('2.2 Deneyim Kriterleri', level=2)
deneyim = [
    ('Minibüs (14+1)', '3 yıl', 'Ticari yolcu taşımacılığı'),
    ('Midibüs (25+1)', '5 yıl', 'Ticari yolcu taşımacılığı'),
]
t = doc.add_table(rows=3, cols=3)
t.style = 'Table Grid'
add_header_row(t, ['Araç Tipi', 'Minimum Deneyim Süresi', 'Ticari Araç Türü'])
for i, row_data in enumerate(deneyim):
    add_table_row(t, i+1, row_data, alt_row=(i%2==0))
set_cell_borders(t)
doc.add_paragraph()

doc.add_heading('Disiplin Şartları:', level=3)
add_bullet(doc, 'Son 3 yılda alkollü araç kullanımından ceza almamış olma (noter onaylı beyan veya trafik kayıt çıktısı).')
add_bullet(doc, 'Son 3 yılda sürücünün kusurundan kaynaklanan ölümlü veya ağır yaralanmalı kaza bulunmaması.')
add_bullet(doc, 'Kamu/özel sektörden ihraç veya disiplin cezası almamış olma.')

doc.add_heading('2.3 Davranış ve Görünüm Standartları', level=2)
doc.add_heading('Üniforma:', level=3)
add_bullet(doc, 'Tüm sürücüler yüklenici tarafından karşılanan üniforma giymekle yükümlüdür.')
add_bullet(doc, 'Üniforma üzerinde yüklenici firmanın logosu bulunacaktır; idare onayına tabidir.')
add_bullet(doc, 'Sürücüler için standart üniforma: uzun kollu gömlek + pantolon veya firma ceketi.')

doc.add_heading('Davranış Kuralları:', level=3)
add_bullet(doc, 'Seyir halinde telefon kullanımı kesinlikle yasaktır (eller serbest kit dahil).')
add_bullet(doc, 'Seyir sırasında sigara, yiyecek veya içecek tüketimi yasaktır.')
add_bullet(doc, 'Yolculara nazik ve saygılı davranmak zorunludur; şikayete konu olan sürücü hakkında yüklenici 5 iş günü içinde yazılı önlem alacaktır.')
add_bullet(doc, 'Güzergah dışı durak yapılması, yolcu almak veya bırakmak için onaysız sapma yapılması yasaktır.')

doc.add_heading('Eğitim ve Gelişim:', level=3)
add_bullet(doc, 'Yılda en az 1 kez yüklenici tarafından sürücülere savunmacı sürüş ve yolcu ilişkileri eğitimi verilecektir.')
add_bullet(doc, 'Eğitim katılım listeleri idareye iletilecektir.')

# ────────────────────────────
# BÖLÜM 3 — GÜZERGAH VE SEFER PLANI
# ────────────────────────────
doc.add_page_break()
add_section_break_line(doc)
h = doc.add_heading('BÖLÜM 3 — GÜZERGAH VE SEFER PLANI', level=1)
h.runs[0].font.color.rgb = RGBColor(0x1F, 0x5C, 0x99)

doc.add_heading('3.1 Hat Tipleri ve Tanımları', level=2)
hat_tipleri = [
    ('Ana Hat', 'Merkez ofis ↔ toplu ulaşım bağlantı noktaları (metro, otobüs terminali)', 'Midibüs (25+1)'),
    ('Bölgesel Hat', 'İlçe bazlı personel toplama güzergahları', 'Minibüs (14+1) veya Midibüs'),
    ('Özel Hat', 'Uzak mesafe veya geç vardiya (gece) personeli için', 'Minibüs (14+1)'),
    ('Okul/Kreş Bağlantı Hattı', 'Personel çocukları için talep edildiğinde', 'Minibüs (14+1)'),
]
t = doc.add_table(rows=5, cols=3)
t.style = 'Table Grid'
add_header_row(t, ['Hat Tipi', 'Açıklama', 'Araç Tipi'])
for i, row_data in enumerate(hat_tipleri):
    add_table_row(t, i+1, row_data, alt_row=(i%2==0))
set_cell_borders(t)
doc.add_paragraph()

add_bullet(doc, 'Hat tanımları, durak listesi ve haritaları sözleşme ekinde yer alacaktır.')
add_bullet(doc, 'Yüklenici, idarenin yazılı talebi halinde 5 iş günü içinde yeni hat veya durak düzenlemesi yapacaktır.')

doc.add_heading('3.2 Sefer Standartları', level=2)
doc.add_heading('Kalkış Saatleri:', level=3)
add_bullet(doc, 'Sabah servisi: toplu kalkış saatinden en az 30 dakika önce ilk durakta hazır bulunulacaktır.')
add_bullet(doc, 'Akşam servisi: vardiya bitiş saatinden en fazla 10 dakika sonra hareket edilecektir.')
add_bullet(doc, 'Geç vardiya servisleri: kalkış saatlerinden en geç 15 dakika önce konuşlandırma tamamlanacaktır.')

doc.add_heading('Güzergah Disiplini:', level=3)
add_bullet(doc, 'Rota sapması: onaylı güzergahtan %15\'ten fazla sapma yasaktır.')
add_bullet(doc, 'Trafikten bağımsız zorunlu sapma durumlarında sürücü merkezi derhal bilgilendirir ve onay alır.')
add_bullet(doc, 'Yetkisiz güzergah değişikliği ceza tablosuna göre yaptırıma tabi tutulur.')

doc.add_heading('Kapasite Kuralları:', level=3)
add_bullet(doc, 'Araçlar kapasitelerinin üzerinde yolcu taşıyamaz.')
add_bullet(doc, 'Ayakta yolcu taşıması kesinlikle yasaktır.')
add_bullet(doc, 'Doluluk oranı anlık olarak GPS sistemi üzerinden izlenecektir.')

doc.add_heading('3.3 Sefer Takip ve Raporlama', level=2)
doc.add_heading('Günlük Sefer Raporu (her seferden sonra sistem otomatik üretir):', level=3)
gunluk_items = ['Araç plakası', 'Sürücü adı ve sicil numarası', 'Güzergah kodu ve hat adı',
                'Planlanan kalkış / gerçekleşen kalkış saati', 'Planlanan varış / gerçekleşen varış saati',
                'Toplam km', 'Yolcu sayısı (biniş / iniş)', 'Olağandışı durum notu (varsa)']
for item in gunluk_items:
    add_bullet(doc, item)

doc.add_heading('Aylık Özet Rapor (her ayın 5. iş günü idareye sunulur):', level=3)
aylik_items = ['Hat bazında sefer gerçekleşme oranı', 'İptal edilen seferler ve gerekçeleri',
               'Ortalama dakiklik oranı', 'Sürücü davranış ihlalleri özeti', 'Yolcu şikayet özeti ve alınan önlemler']
for item in aylik_items:
    add_bullet(doc, item)

# ────────────────────────────
# BÖLÜM 4 — BAKIM VE TEMİZLİK STANDARTLARI
# ────────────────────────────
doc.add_page_break()
add_section_break_line(doc)
h = doc.add_heading('BÖLÜM 4 — BAKIM VE TEMİZLİK STANDARTLARI', level=1)
h.runs[0].font.color.rgb = RGBColor(0x1F, 0x5C, 0x99)

doc.add_heading('4.1 Periyodik Bakım Planı', level=2)
bakim_data = [
    ('Yağ değişimi ve genel kontrol', 'Her 10.000 km', 'Yetkili servis faturası + bakım formu'),
    ('Fren sistemi kontrolü', '6 ayda bir', 'Yetkili servis raporu'),
    ('Klima filtresi değişimi', '3 ayda bir', 'İç bakım kaydı'),
    ('Tam araç muayenesi', 'Yılda bir', 'Bağımsız yetkili servis raporu'),
    ('Lastik kontrolü ve rotasyon', '15.000 km\'de bir', 'İç bakım kaydı'),
    ('Akü ve elektrik sistemi kontrolü', '6 ayda bir', 'İç bakım kaydı'),
    ('Şanzıman ve diferansiyel yağ kontrolü', '40.000 km\'de bir', 'Yetkili servis faturası'),
]
t = doc.add_table(rows=len(bakim_data)+1, cols=3)
t.style = 'Table Grid'
add_header_row(t, ['Bakım Türü', 'Periyot', 'Belgeleme'])
for i, row_data in enumerate(bakim_data):
    add_table_row(t, i+1, row_data, alt_row=(i%2==0))
set_cell_borders(t)
doc.add_paragraph()

doc.add_heading('Lastik Standartları:', level=3)
add_bullet(doc, 'Lastik profil derinliği: minimum 3 mm (genel dönem).')
add_bullet(doc, 'Kış lastiği döneminde (1 Kasım – 31 Mart): minimum 4 mm profil derinliği ve M+S / kış lastiği sertifikası.')
add_bullet(doc, 'Yıpranmış veya hasarlı lastikle sefer yasaktır; sürücü kontrol yükümlülüğü vardır.')

doc.add_heading('4.2 Temizlik Standartları', level=2)
doc.add_heading('Günlük İç Temizlik (Her Seferden Önce):', level=3)
ic_temizlik = ['Koltuklar — görünür kir ve lekelerden arındırılmış', 'Zemin — süpürme ve nemli silme',
               'Camlar (iç yüzey) — lekelere karşı silinmiş', 'Tutunma kolları ve emniyet kemeri tokmaları — silinmiş',
               'Çöp bidonları — boşaltılmış']
for item in ic_temizlik:
    add_bullet(doc, item)

doc.add_heading('Haftalık Derin Temizlik:', level=3)
haftalik = ['Koltuk kılıfları değiştirilecek ve yıkandıktan sonra tekrar takılacaktır.',
            'Araç içi tüm yüzeyler dezenfektanla silinecektir.',
            'Hava filtreleri kontrol edilecektir.', 'Bagaj bölümleri (varsa) temizlenecektir.']
for item in haftalik:
    add_bullet(doc, item)

doc.add_heading('Dış Yıkama:', level=3)
add_bullet(doc, 'Haftada en az 3 (üç) kez dış yıkama yapılacaktır.')
add_bullet(doc, 'Görünür kirlenme (çamur, kuş pisliği vb.) halinde sefer öncesi yıkama zorunludur.')

doc.add_heading('Dezenfeksiyon:', level=3)
add_bullet(doc, 'Haftalık tam araç dezenfeksiyonu yapılacaktır.')
add_bullet(doc, 'Salgın hastalık veya idare talebi halinde sıklık artırılacaktır.')
add_bullet(doc, 'Yalnızca Tarım ve Orman Bakanlığı\'nca onaylı dezenfektan ürünleri kullanılacaktır.')

# ────────────────────────────
# BÖLÜM 5 — PERFORMANS GÖSTERGELERİ (KPI)
# ────────────────────────────
doc.add_page_break()
add_section_break_line(doc)
h = doc.add_heading('BÖLÜM 5 — PERFORMANS GÖSTERGELERİ (KPI)', level=1)
h.runs[0].font.color.rgb = RGBColor(0x1F, 0x5C, 0x99)

doc.add_heading('5.1 KPI Tablosu', level=2)
add_body(doc, 'Hizmet kalitesi aşağıdaki ölçülebilir performans göstergeleri ile izlenecektir. Yüklenici, aylık raporda her gösterge için gerçekleşme değerini sunacaktır.')
doc.add_paragraph()

kpi_data = [
    ('KPI-01', 'Sefer Gerçekleşme Oranı', '%98', '<%96', '<%95 → yazılı uyarı; <%90 → sözleşme cezası'),
    ('KPI-02', 'Dakiklik Oranı (±5 dk tolerans)', '%95', '<%92', '<%90 → hakediş kesintisi'),
    ('KPI-03', 'Araç Arızası Sonrası İkame Süresi', '≤ 30 dakika', '31–45 dk', '> 45 dk → sefer başına ceza'),
    ('KPI-04', 'Yolcu Şikayeti Çözüm Süresi', '≤ 24 saat', '25–48 saat', '> 48 saat → yazılı uyarı'),
    ('KPI-05', 'GPS Veri İletim Sürekliliği', '%99', '<%98', '<%97 → yazılı uyarı; <%95 → teknik denetim'),
    ('KPI-06', 'Aylık Sefer Raporu Teslim Süresi', '5. iş günü', '6.–7. iş günü', '8. iş gününden sonra → uyarı'),
    ('KPI-07', 'Araç Temizlik Uygunluk Oranı', '%100', '<%98', '<%95 → tutanak + ceza'),
    ('KPI-08', 'Sürücü Üniforması Uyum Oranı', '%100', '<%99', '<%97 → tutanak'),
    ('KPI-09', 'Yolcu Memnuniyet Skoru', '≥ 4,0 / 5,0', '3,5–3,9', '< 3,5 → iyileştirme planı talebi'),
    ('KPI-10', 'Kapasite Aşımı Vakası', '0', '1–2 vaka/ay', '≥ 3 vaka/ay → ihtar'),
    ('KPI-11', 'Güzergah Sapma Vakası', '0', '1–2 vaka/ay', '≥ 3 vaka/ay → yazılı uyarı + hakediş kesintisi'),
]
t = doc.add_table(rows=len(kpi_data)+1, cols=5)
t.style = 'Table Grid'
add_header_row(t, ['#', 'Performans Göstergesi', 'Hedef Değer', 'Uyarı Eşiği', 'Ceza / Yaptırım Eşiği'])
for i, row_data in enumerate(kpi_data):
    add_table_row(t, i+1, row_data, bold_first=True, alt_row=(i%2==0))
set_cell_borders(t)
doc.add_paragraph()

doc.add_heading('5.2 Ölçüm ve Değerlendirme Yöntemi', level=2)
doc.add_heading('Değerlendirme Dönemleri:', level=3)
add_bullet(doc, 'Aylık: KPI-01\'den KPI-08\'e kadar tüm göstergeler')
add_bullet(doc, 'Üç aylık: KPI-09 yolcu memnuniyet skoru')
add_bullet(doc, 'Yıllık: Genel performans değerlendirme raporu')

doc.add_heading('Ceza Mekanizması:', level=3)
add_bullet(doc, 'Yazılı uyarılar yüklenicinin kayıtlı elektronik posta adresine bildirilir.')
add_bullet(doc, 'Hakediş kesintileri; idare tarafından aylık ödeme tablosuna yansıtılır.')
add_bullet(doc, 'Ardı ardına 3 ay ceza eşiğinin altında kalan yüklenici hakkında sözleşme feshi prosedürü başlatılabilir.')

# ────────────────────────────
# BÖLÜM 6 — ACİL DURUM PROSEDÜRLERİ
# ────────────────────────────
doc.add_page_break()
add_section_break_line(doc)
h = doc.add_heading('BÖLÜM 6 — ACİL DURUM PROSEDÜRLERİ', level=1)
h.runs[0].font.color.rgb = RGBColor(0x1F, 0x5C, 0x99)

doc.add_heading('6.1 Araç Arızası Prosedürü', level=2)
add_numbered(doc, 'Sürücü güvenli bir alana çeker ve araç ikaz lambalarını aktif eder.')
add_numbered(doc, 'Yolcular araçtan güvenli biçimde indirilir.')
add_numbered(doc, 'Sürücü yüklenici merkez operatörünü derhal (5 dakika içinde) arar.')
add_numbered(doc, 'Operatör, en yakın yedek aracı ve sürücüyü göreve yönlendirir.')

doc.add_heading('Yedek Araç Devreye Alma:', level=3)
add_bullet(doc, 'Yedek araç, çağrı alındıktan itibaren 30 dakika içinde yola çıkar.')
add_bullet(doc, 'Yolcular yedek araçla güzergah tamamlanır veya idare onayı ile alternatif ulaşım sağlanır.')
add_bullet(doc, 'Alternatif ulaşım masrafları yüklenici tarafından karşılanır.')
add_bullet(doc, 'Arıza tutanağı dijital ortamda 2 saat içinde idareye iletilir.')

doc.add_heading('6.2 Trafik Kazası Prosedürü', level=2)
add_numbered(doc, 'Sürücü mümkünse güvenli alana geçer, motor kapatılır.')
add_numbered(doc, 'Yaralı varsa 112 Acil aranır; yolcuların güvenliği önceliklidir.')
add_numbered(doc, 'İdare acil hattı eş zamanlı olarak aranır.')
add_numbered(doc, 'Sürücü kaza mahallini terk edemez; yetkili makamların (trafik polisi) onayı beklenir.')
add_note(doc, 'Kaza tutanağı dijital ortamda 2 saat içinde idareye iletilir. Fotoğraf ekleri (en az 3 farklı açı) zorunludur.')

doc.add_heading('6.3 Hava Koşulları ve Doğal Afet Prosedürü', level=2)
add_bullet(doc, 'Kar ve buzlanma uyarısı halinde kış lastiği takıntısı zorunludur.')
add_bullet(doc, 'Görüş mesafesi 50 metrenin altına düştüğünde sefer durdurulur.')
add_bullet(doc, 'AFAD veya yerel yönetimin ilan ettiği afet durumunda tüm seferler durdurulur.')
add_bullet(doc, 'İdare koordinasyonuyla yeniden başlama tarihi belirlenir.')

doc.add_heading('6.4 Sürücü Rahatsızlığı veya Görev Yapamaması Prosedürü', level=2)
add_numbered(doc, 'Sürücü güvenli alana çeker ve motoru kapatır.')
add_numbered(doc, 'Gerekirse 112 aranır; merkeze bildirim yapılır ve yedek sürücü talep edilir.')
add_numbered(doc, 'Yedek sürücü gelene kadar araç ve yolcular güvende bekler.')
add_note(doc, 'Yüklenici yedek sürücü havuzu tutacaktır; her aktif sürücü için en az %10 kapasitede yedek sürücü hazır olacaktır.')

# ────────────────────────────
# BÖLÜM 7 — YASAL DAYANAK VE MEVZUAT UYUMU
# ────────────────────────────
doc.add_page_break()
add_section_break_line(doc)
h = doc.add_heading('BÖLÜM 7 — YASAL DAYANAK VE MEVZUAT UYUMU', level=1)
h.runs[0].font.color.rgb = RGBColor(0x1F, 0x5C, 0x99)

doc.add_heading('7.1 Amaç ve Kapsam', level=2)
add_body(doc, 'Bu bölüm, personel servisi hizmet alımına ilişkin teknik şartnamenin hazırlanmasında, uygulanmasında ve denetlenmesinde esas alınan yasal çerçeveyi belirler. Yüklenici, hizmetin ifası süresince aşağıda sayılan mevzuat hükümlerine tam ve eksiksiz uymakla yükümlüdür.')

doc.add_heading('7.2 Temel Tanımlar', level=2)
hukuki_tanımlar = [
    ('İdare', 'İhaleyi yürüten ve sözleşmeyi imzalayan kamu kurumu'),
    ('Yüklenici', 'İhaleyi kazanan ve hizmet alımı sözleşmesini imzalayan gerçek veya tüzel kişi'),
    ('Sefer', 'Belirlenen güzergah üzerinde, öngörülen saatlerde gerçekleştirilen personel taşıma hareketi'),
    ('Güzergah', 'Sefer kapsamında araçların izleyeceği onaylı rota; durak noktaları dahil'),
    ('Aktif Filo', 'Sözleşme kapsamında fiilen hizmet veren araç sayısı'),
    ('Yardımcı Personel', 'Sürücü dışında araçta görev yapan hostes veya güvenlik görevlisi'),
    ('KPI', 'Hizmet kalite göstergesi (Key Performance Indicator); Bölüm 5\'te tanımlı'),
    ('İhlal Tespit Tutanağı', 'İdarece düzenlenen, yüklenicinin yükümlülüğünü yerine getirmediğini belgeleyen resmi tutanak'),
]
t = doc.add_table(rows=len(hukuki_tanımlar)+1, cols=2)
t.style = 'Table Grid'
add_header_row(t, ['Kavram', 'Tanım'])
for i, row_data in enumerate(hukuki_tanımlar):
    add_table_row(t, i+1, row_data, bold_first=True, alt_row=(i%2==0))
set_cell_borders(t)
doc.add_paragraph()

doc.add_heading('7.3 İlgili Mevzuat Tablosu', level=2)
mevzuat = [
    ('1', 'Kamu İhale Kanunu', '4734 Sayılı Kanun', 'İhale usul ve esasları, isteklilerin nitelikleri'),
    ('2', 'Kamu İhale Sözleşmeleri Kanunu', '4735 Sayılı Kanun', 'Sözleşme türleri, fesih, tazminat hükümleri'),
    ('3', 'Karayolu Taşıma Kanunu', '4925 Sayılı Kanun', 'Karayoluyla yolcu taşımacılığı, yetki belgeleri'),
    ('4', 'Karayolu Taşıma Yönetmeliği', 'RG 11.06.2009/27255', 'D tipi yetki belgesi gereklilikleri'),
    ('5', 'Karayolları Trafik Kanunu', '2918 Sayılı Kanun', 'Trafik kuralları, araç teknik koşulları'),
    ('6', 'Karayolları Trafik Yönetmeliği', 'RG 18.07.1997/23053', 'Araç muayenesi, ehliyet sınıfları'),
    ('7', 'İş Sağlığı ve Güvenliği Kanunu', '6331 Sayılı Kanun', 'Çalışan sağlığı, risk değerlendirmesi'),
    ('8', 'Kişisel Verilerin Korunması Kanunu', '6698 Sayılı Kanun', 'Araç kamerası kayıtları, kişisel veri koruması'),
    ('9', 'SGK Kanunu', '5510 Sayılı Kanun', 'Sürücü sigorta primleri, iş kazası bildirimi'),
    ('10', 'Türk Borçlar Kanunu', '6098 Sayılı Kanun', 'Hizmet sözleşmesi, sorumluluk, tazminat'),
    ('11', 'İş Kanunu', '4857 Sayılı Kanun', 'Çalışma süreleri, fazla mesai, işçi hakları'),
    ('12', 'Motorlu Taşıtlar Vergisi Kanunu', '197 Sayılı Kanun', 'Araç vergi yükümlülükleri'),
    ('13', 'KTY — Sürücü Çalışma Süreleri', 'Madde 75–82', 'Sürüş süreleri, dinlenme zorunlulukları'),
    ('14', 'Kamu İhale Genel Tebliği', 'KİK (Güncel)', 'Teknik şartname hazırlama usulü'),
    ('15', 'Sigortacılık Kanunu', '5684 Sayılı Kanun', 'Sigorta poliçeleri, tazminat'),
    ('16', 'İcra ve İflas Kanunu', '2004 Sayılı Kanun', 'İflas, konkordato ve tasfiye hükümleri'),
    ('17', 'Arabuluculuk Kanunu', '6325 Sayılı Kanun', 'İhtiyari arabuluculuk usulü'),
    ('18', 'İdari Yargılama Usulü Kanunu', '2577 Sayılı Kanun', 'Kamu sözleşmelerinde idari yargı yetkisi'),
]
t = doc.add_table(rows=len(mevzuat)+1, cols=4)
t.style = 'Table Grid'
add_header_row(t, ['No', 'Mevzuat', 'Kanun/Yönetmelik No', 'Düzenleme Konusu'])
for i, row_data in enumerate(mevzuat):
    add_table_row(t, i+1, row_data, alt_row=(i%2==0))
set_cell_borders(t)
doc.add_paragraph()

doc.add_heading('7.4 Lisans ve Yetki Belgeleri', level=2)
doc.add_heading('7.4.1 Kurumsal Belgeler', level=3)
add_bullet(doc, 'Ticaret Sicil Gazetesi ve Sicil Kayıt Sureti (güncel, son 6 ay içinde alınmış)')
add_bullet(doc, 'İmza Sirküleri (noter tasdikli, güncel)')
add_bullet(doc, 'Vergi Levhası ve vergi borcu yoktur yazısı (son 1 ay içinde alınmış)')
add_bullet(doc, 'SGK Borcu Yoktur Yazısı (son 1 ay içinde alınmış, e-Devlet onaylı)')

doc.add_heading('7.4.2 Taşımacılık Yetki Belgeleri', level=3)
add_bullet(doc, 'D2 Yetki Belgesi — Ücretli yolcu taşımacılığı; Ulaştırma ve Altyapı Bakanlığı onaylı.')
add_bullet(doc, 'Araç sayısı kadar D2 yetki belgesi kapsamında araç kayıt belgesi.')

doc.add_heading('7.4.3 Sürücülere Ait Belgeler', level=3)
add_bullet(doc, 'Sürücü Belgesi — En az D sınıfı (9 ve üzeri koltuklu araçlar için).')
add_bullet(doc, 'Psikoteknik Değerlendirme Belgesi — 4925 sk Madde 26 ve KTY Madde 72 uyarınca zorunlu.')
add_bullet(doc, 'Mesleki Yeterlilik Belgesi (SRC-3) — KTY Madde 74 kapsamında yolcu taşımacılığı sürücüleri için zorunludur. (NOT: Yolcu taşımacılığında SRC-3 geçerli belgedir; SRC-2 yük taşımacılığına özgü olup bu şartname kapsamında kabul edilmez.)')
add_bullet(doc, 'Sağlık Raporu — Periyodik sürücü sağlık muayenesi belgesi, yetkili sağlık kurumu onaylı, en fazla 1 yıl öncesine ait.')

# ────────────────────────────
# BÖLÜM 8 — SİGORTA VE SORUMLULUK
# ────────────────────────────
doc.add_page_break()
add_section_break_line(doc)
h = doc.add_heading('BÖLÜM 8 — SİGORTA VE SORUMLULUK HÜKÜMLERİ', level=1)
h.runs[0].font.color.rgb = RGBColor(0x1F, 0x5C, 0x99)

doc.add_heading('8.1 Hukuki Dayanak', level=2)
add_body(doc, '2918 Sayılı Karayolları Trafik Kanunu, Zorunlu Mali Sorumluluk Sigortası Tarife ve Talimatları, 6098 Sayılı Türk Borçlar Kanunu\'nun 49–76. maddeleri ve 5684 Sayılı Sigortacılık Kanunu hükümleri esas alınmıştır.')

doc.add_heading('8.2 Zorunlu Sigorta Türleri ve Asgari Teminat Tutarları', level=2)
sigorta_data = [
    ('Zorunlu Mali Sorumluluk Sigortası (Trafik Sigortası)', 'Hazine ve Maliye Bakanlığı\'nın yıllık yasal asgari limitleri', 'Üçüncü şahıslara maddi ve bedeni zararlar', '2918 sk Madde 91'),
    ('Ferdi Kaza Sigortası (Koltuk Sigortası)', 'Koltuk başına asgari 750.000 TL (ölüm ve sürekli sakatlık)', 'Araç içindeki yolcuların kaza zararları', '2918 sk Madde 93'),
    ('Kasko Sigortası', 'Aracın güncel piyasa değerinin %100\'ü (azami %5 muafiyetli)', 'Araç hasarı (çarpışma, yangın, hırsızlık, doğal afet)', 'Sigortacılık Kanunu (5684 sk)'),
    ('İşveren Mali Sorumluluk Sigortası', 'Her çalışan başına asgari 3.000.000 TL', 'Sürücü ve personelin iş kazası ve meslek hastalıkları', '6331 sk Madde 21; TBK Madde 417'),
    ('Mesleki Sorumluluk Sigortası', 'Asgari 5.000.000 TL (poliçe dönemi toplam)', 'Hizmet kusuru, ihmal, yanlış uygulama', '4925 sk; KTY'),
    ('Üçüncü Şahıs Sorumluluk Sigortası', 'Asgari 10.000.000 TL (olay başına)', 'Yolcular dışındaki üçüncü kişilere zararlar', 'TBK Madde 49'),
]
t = doc.add_table(rows=len(sigorta_data)+1, cols=4)
t.style = 'Table Grid'
add_header_row(t, ['Sigorta Türü', 'Asgari Teminat', 'Kapsam', 'Dayanak'])
for i, row_data in enumerate(sigorta_data):
    add_table_row(t, i+1, row_data, alt_row=(i%2==0))
set_cell_borders(t)
doc.add_paragraph()

add_note(doc, 'Enflasyon Güncelleme Mekanizması: TL cinsinden sabit teminat tutarları, her yılın Ocak ayında bir önceki yılın TÜFE (12 aylık birikimli) oranında otomatik olarak güncellenir. Güncellenmiş tutarlar idare tarafından yükleniciye yazılı olarak bildirilir. Yüklenici, yenilenen teminat tutarlarını yansıtan güncel poliçeleri bildirim tarihinden itibaren 20 iş günü içinde idareye sunar.')

doc.add_heading('8.3 Poliçe İbrazı ve Yenileme Yükümlülüğü', level=2)
add_bullet(doc, 'Sigorta poliçeleri, sözleşme imzalanmadan önce idareye teslim edilir.')
add_bullet(doc, 'Yıllık yenilemelerde, poliçe bitiş tarihinden en az 10 iş günü önce yenilenmiş poliçe idareye teslim edilir.')
add_bullet(doc, 'Poliçe teslimatında gecikme halinde, gecikilen her gün için günlük hizmet bedelinin %0,5\'i oranında ceza uygulanır.')
add_bullet(doc, 'Sigorta şirketinin iflası veya lisans iptali halinde yüklenici 5 iş günü içinde yeni poliçe yaptırır.')

doc.add_heading('8.4 Sorumluluk Hükümleri', level=2)
add_body(doc, 'Hizmet kapsamındaki her türlü kaza, hasar, yaralanma ve ölüm olayında birincil ve tam sorumluluk yükleniciye aittir. İdare, 6098 Sayılı TBK\'nın 66. maddesi kapsamında adam çalıştıran sıfatıyla yüklenici çalışanlarının eylemlerinden dolayı üçüncü kişilere karşı sorumlu tutulmaz.')
doc.add_paragraph()
add_note(doc, 'Kusursuz sorumluluk (TBK Madde 71): Yüklenici, araçların işletilmesinden kaynaklanan zararlardan kusuru olmaksızın da sorumludur. İdarenin rücu hakkı saklıdır (TBK Madde 500).')

# ────────────────────────────
# BÖLÜM 9 — KVKK
# ────────────────────────────
doc.add_page_break()
add_section_break_line(doc)
h = doc.add_heading('BÖLÜM 9 — KİŞİSEL VERİLERİN KORUNMASI (KVKK)', level=1)
h.runs[0].font.color.rgb = RGBColor(0x1F, 0x5C, 0x99)

doc.add_heading('9.1 Hukuki Dayanak', level=2)
add_body(doc, '6698 Sayılı Kişisel Verilerin Korunması Kanunu (KVKK), Kişisel Verilerin Silinmesi, Yok Edilmesi veya Anonim Hale Getirilmesi Hakkında Yönetmelik (RG: 28.10.2017/30224) ve Kişisel Veri Güvenliği Rehberi (KVKK Kurulu) hükümleri kapsamında düzenlenmiştir.')

doc.add_heading('9.2 Veri Sorumlusu ve Veri İşleyen Sıfatı', level=2)
add_bullet(doc, 'Veri Sorumlusu: İhaleyi yürüten kamu kurumu (İdare), yolcu verilerine ilişkin veri sorumlusu sıfatını haizdir.')
add_bullet(doc, 'Veri İşleyen: Yüklenici, yalnızca hizmetin ifası kapsamında ve İdare\'nin yazılı talimatı doğrultusunda kişisel verileri işler.')
add_bullet(doc, 'Yüklenici, veri işleme faaliyetlerine başlamadan önce İdare ile Veri İşleme Sözleşmesi (DPA) akdeder.')

doc.add_heading('9.3 Kişisel Veri Kategorileri ve İşleme Amaçları', level=2)
kvkk_data = [
    ('Kimlik verisi', 'Yolcu adı, sicil no', 'Yolcu kaydı, sefer planlaması', 'Madde 5/2-c (sözleşmenin ifası)'),
    ('Konum verisi', 'GPS konumu, sefer güzergahı', 'Güzergah optimizasyonu, güvenlik', 'Madde 5/2-f (meşru menfaat)'),
    ('Görsel veri', 'Araç içi kamera kayıtları', 'Güvenlik, kaza incelemesi', 'Madde 5/2-f; Madde 6 kapsamı dışı'),
    ('Biyometrik veri', 'Yüz tanıma (kullanılıyorsa)', '—', 'Açık rıza zorunlu (Madde 6/3)'),
    ('Sağlık verisi', 'Sürücü sağlık raporu', 'İSG yükümlülüğü', 'Madde 6/3 (açık rıza) + 6/2-ç'),
]
t = doc.add_table(rows=len(kvkk_data)+1, cols=4)
t.style = 'Table Grid'
add_header_row(t, ['Veri Kategorisi', 'Örnekler', 'İşleme Amacı', 'Hukuki Dayanak (KVKK)'])
for i, row_data in enumerate(kvkk_data):
    add_table_row(t, i+1, row_data, alt_row=(i%2==0))
set_cell_borders(t)
doc.add_paragraph()

doc.add_heading('9.4 Veri Saklama Süreleri ve İmha', level=2)
saklama_data = [
    ('Araç içi kamera kayıtları', '30 takvim günü', 'Güvenli silme (overwrite/şifreleme)', 'KVKK Md. 7'),
    ('Aktif soruşturma kapsamındaki kamera kaydı', 'Soruşturma + 30 gün', 'Adli bilişim uzmanı gözetiminde', '5271 sk CMK; KVKK Md. 28'),
    ('GPS/konum verileri', '90 takvim günü', 'Güvenli silme', 'KVKK Md. 7'),
    ('Yolcu sefer kayıtları', '5 yıl', 'Fiziksel imha veya güvenli silme', '4735 sk Madde 36'),
    ('Sürücü özlük dosyaları', 'İş ilişkisi + 10 yıl', 'Gizlilik içinde imha', '4857 sk; KVKK Md. 7'),
    ('Elektronik haberleşme kayıtları', '1 yıl', 'Güvenli silme', '5651 sk Madde 5'),
]
t = doc.add_table(rows=len(saklama_data)+1, cols=4)
t.style = 'Table Grid'
add_header_row(t, ['Veri Türü', 'Saklama Süresi', 'İmha Yöntemi', 'Dayanak'])
for i, row_data in enumerate(saklama_data):
    add_table_row(t, i+1, row_data, alt_row=(i%2==0))
set_cell_borders(t)
doc.add_paragraph()

doc.add_heading('9.5 Veri Güvenliği Teknik Tedbirleri', level=2)
add_bullet(doc, 'Kamera sistemlerinde şifreli depolama (AES-256 veya eşdeğeri)')
add_bullet(doc, 'Verilere erişim yetkisi rol tabanlı erişim kontrolü (RBAC) ile sınırlandırılır.')
add_bullet(doc, 'GPS ve sefer verilerinin aktarımında uçtan uca şifreleme (TLS 1.2 veya üzeri)')
add_bullet(doc, 'Ağ altyapısında güvenlik duvarı ve saldırı tespit sistemi (IDS/IPS)')
add_bullet(doc, 'Yetkisiz erişim girişimlerine karşı erişim loglarının tutulması ve 90 gün saklanması')
add_bullet(doc, 'Veri işleme sistemlerine yönelik yıllık sızma testi (penetrasyon testi) yapılması')

doc.add_heading('9.6 Veri İhlali Bildirim Yükümlülüğü', level=2)
veri_ihlal = [
    ('24 saat içinde', 'İdare', 'Ön bildirim (ihlal özeti, etkilenen veri kategorisi)'),
    ('72 saat içinde', 'KVKK Kurulu (kvkk.gov.tr)', 'Tam bildirim formu (KVKK Madde 12/5)'),
    ('Makul sürede', 'Etkilenen ilgili kişiler', 'Aydınlatma bildirimi'),
]
t = doc.add_table(rows=len(veri_ihlal)+1, cols=3)
t.style = 'Table Grid'
add_header_row(t, ['Süre', 'Bildirim Yapılacak Makam', 'İçerik'])
for i, row_data in enumerate(veri_ihlal):
    add_table_row(t, i+1, row_data, alt_row=(i%2==0))
set_cell_borders(t)

# ────────────────────────────
# BÖLÜM 10 — SÖZLEŞME İHLALİ VE YAPTIRIMLAR
# ────────────────────────────
doc.add_page_break()
add_section_break_line(doc)
h = doc.add_heading('BÖLÜM 10 — SÖZLEŞME İHLALİ VE YAPTIRIMLAR', level=1)
h.runs[0].font.color.rgb = RGBColor(0x1F, 0x5C, 0x99)

doc.add_heading('10.1 Hukuki Dayanak', level=2)
add_body(doc, '4735 Sayılı Kamu İhale Sözleşmeleri Kanunu\'nun 20–22. maddeleri, 4734 Sayılı Kamu İhale Kanunu\'nun ilgili hükümleri ve 6098 Sayılı Türk Borçlar Kanunu\'nun sözleşme ihlali ve tazminata ilişkin hükümleri kapsamında düzenlenmiştir.')

doc.add_heading('10.2 Cezai Hükümler (Götürü Ceza Tablosu)', level=2)
ceza_data = [
    ('1', 'Seferin gerekçesiz iptal edilmesi (24 saat önceden bildirim yapılmadan)', 'Günlük hizmet bedelinin %3\'ü', 'Her vaka', 'Aylık hizmet bedelinin %15\'i'),
    ('2', 'Seferin gecikmeli gerçekleşmesi (15 dk üzeri)', 'Aylık hizmet bedelinin %0,2\'si/olay', 'Her vaka', 'Günlük hizmet bedelinin %10\'u'),
    ('3', 'Araç yaş limitinin aşılması (Bölüm 1 ihlali)', 'Hizmet bedelinin %5\'i/gün', 'Tespetten araç değişimine kadar', '—'),
    ('4', 'Güzergah dışına çıkılması (onaysız)', 'Aylık hizmet bedelinin %0,5\'i/olay', 'Her vaka', '—'),
    ('5', 'GPS/takip sisteminin kapalı olması', 'Aylık hizmet bedelinin %0,3\'ü/saat', '2 saat kesintisizin üzerindeki her saat', '—'),
    ('6', 'Sürücü belgesinin geçersiz/süresi dolmuş olması', 'Hizmet bedelinin %10\'u + derhal araç değişimi', 'Tespit başına', '—'),
    ('7', 'SRC-3 veya psikoteknik belgesinin eksik/geçersiz olması', 'Hizmet bedelinin %7\'si + sürücü değişimi', 'Tespit başına', '—'),
    ('8', 'Araç periyodik muayenesinin yapılmamış olması', 'Hizmet bedelinin %5\'i + araç hizmet dışı', 'Tespit başına', '—'),
    ('9', 'Araç içi kamera kaydının silinmesi (talep/soruşturma öncesinde)', '10.000 TL', 'Her vaka', '—'),
    ('10', 'KVKK ihlali (yetkisiz veri paylaşımı, bildirim yapılmaması)', '20.000 TL', 'Her vaka', '—'),
    ('11', 'İSG eğitimi belgesinin ibraz edilmemesi', '5.000 TL', 'Her eksik sürücü/ay', '—'),
    ('12', 'İş kazasının yasal sürede bildirilmemesi (SGK, Bakanlık)', '15.000 TL', 'Her vaka', '—'),
    ('13', 'Sigorta poliçesinin süresi dolduğu halde yenilenmemesi', 'Günlük hizmet bedelinin %1\'i', 'Gecikilen her gün', '—'),
    ('14', 'KPI değerlerinin art arda 3 ay tutturulamaması', 'Sözleşme fesih hakkı doğar', '—', '—'),
    ('15', 'Yılda 3 kez aynı türde ihlal tekrarı', 'Ceza iki katı uygulanır', 'Her yinelenen ihlal', '—'),
]
t = doc.add_table(rows=len(ceza_data)+1, cols=5)
t.style = 'Table Grid'
add_header_row(t, ['No', 'İhlal Türü', 'Ceza Miktarı', 'Uygulama Birimi', 'Azami Kümülatif Sınır'])
for i, row_data in enumerate(ceza_data):
    add_table_row(t, i+1, row_data, alt_row=(i%2==0))
set_cell_borders(t)
doc.add_paragraph()

add_note(doc, 'Ceza Tutarlarının Güncellenmesi: Tablodaki TL bazlı sabit ceza tutarları (satır 9, 10, 11, 12), her yılın Ocak ayında bir önceki yılın TÜFE oranında güncellenir. Güncel tutarlar, idare tarafından sözleşmenin birinci yılı başlangıcında taraflara duyurulur. "Hizmet bedeli oranı" olarak tanımlanan cezalar enflasyondan otomatik olarak korunmakta olup güncellemeye gerek duyulmaz.')

doc.add_heading('10.3 Ceza Uygulama Usulü', level=2)
add_bullet(doc, 'Cezalar, idarece düzenlenen İhlal Tespit Tutanağı ile belgelenir.')
add_bullet(doc, 'Yükleniciye tutanağın tebliğinden itibaren 5 iş günü itiraz süresi tanınır.')
add_bullet(doc, 'İtirazın reddi veya itiraz yapılmaması halinde ceza, aynı ay içindeki hakedişten veya teminattan mahsup edilir.')
add_bullet(doc, 'Aylık toplam ceza miktarının aylık hizmet bedelinin %25\'ini aşması halinde idare, sözleşmeyi ihtar gereksinimi olmaksızın fesheder.')

doc.add_heading('10.4 Sözleşme Feshi Koşulları', level=2)
doc.add_heading('10.4.1 İdarenin Tazminatsız Fesih Hakkı (4735 sk Madde 20):', level=3)
add_bullet(doc, 'Yüklenicinin iflası, konkordato ilan etmesi veya tasfiyeye girmesi.')
add_bullet(doc, 'Mesleki yetki belgesinin (D2) Ulaştırma Bakanlığı tarafından iptali veya askıya alınması.')
add_bullet(doc, 'Ağır iş güvenliği ihlali sonucu meydana gelen ölümlü iş kazasında yüklenicinin kusurlu olması.')
add_bullet(doc, 'Bir takvim yılı içinde 5\'ten fazla gerekçesiz sefer iptali.')
add_bullet(doc, 'KVKK Kurulu tarafından idari para cezasına çarptırılma veya veri işleme faaliyetinin durdurulması.')
add_bullet(doc, 'Yüklenicinin rüşvet veya yolsuzluk suçundan mahkûm olması (4734 sk Madde 17).')

doc.add_heading('10.4.2 Yüklenicinin Haklı Fesih Hakkı:', level=3)
add_bullet(doc, 'İdarenin hakediş ödemelerini 30 günü aşan süreyle geciktirmesi.')
add_bullet(doc, 'İdare tarafından sözleşme kapsamı dışı iş talep edilmesi.')
add_bullet(doc, 'Mücbir sebep hallerinin (doğal afet, savaş, pandemi) 30 günü aşması.')

# ────────────────────────────
# BÖLÜM 11 — UYUŞMAZLIK ÇÖZÜMÜ
# ────────────────────────────
doc.add_page_break()
add_section_break_line(doc)
h = doc.add_heading('BÖLÜM 11 — UYUŞMAZLIK ÇÖZÜMÜ', level=1)
h.runs[0].font.color.rgb = RGBColor(0x1F, 0x5C, 0x99)

doc.add_heading('11.1 Hukuki Dayanak', level=2)
add_body(doc, '4735 Sayılı Kamu İhale Sözleşmeleri Kanunu Madde 20, 2577 Sayılı İdari Yargılama Usulü Kanunu, 6325 Sayılı Hukuk Uyuşmazlıklarında Arabuluculuk Kanunu ve 4686 Sayılı Milletlerarası Tahkim Kanunu çerçevesinde düzenlenmiştir.')

doc.add_heading('11.2 Uyuşmazlık Çözüm Aşamaları', level=2)
doc.add_heading('11.2.1 Birinci Kademe — İdari Müzakere:', level=3)
add_bullet(doc, 'Uyuşmazlık doğduğu anda taraf, diğer tarafa yazılı bildirim gönderir.')
add_bullet(doc, 'Bildirimi alan taraf, tebellüğden itibaren 10 iş günü içinde yazılı yanıt vermek zorundadır.')
add_bullet(doc, 'Taraflar, yazışmadan itibaren 15 iş günü içinde müzakere toplantısı düzenler.')
add_bullet(doc, 'Müzakere sonucu tutanağa bağlanır ve her iki tarafın yetkili imzası alınır.')

doc.add_heading('11.2.2 İkinci Kademe — Arabuluculuk:', level=3)
add_bullet(doc, 'İdari müzakereden sonuç alınamazsa 30 takvim günü içinde arabuluculuk başvurusu yapılabilir.')
add_bullet(doc, 'Arabuluculuk, 6325 Sayılı Kanun hükümleri çerçevesinde yürütülür.')
add_bullet(doc, 'Arabuluculuk süreci en fazla 60 takvim günü içinde tamamlanır.')

doc.add_heading('11.2.3 Üçüncü Kademe — Yargı Yolu:', level=3)
add_bullet(doc, 'Arabuluculuktan sonuç alınamazsa yetkili idare mahkemesine başvurulur.')
add_bullet(doc, 'Yetki: Sözleşmenin imzalandığı yer İdare Mahkemesi, 2577 sk Madde 32–33 uyarınca yetkilidir.')
add_bullet(doc, 'Uygulanacak hukuk Türk Hukuku\'dur; dil Türkçedir.')

doc.add_heading('11.3 Zamanaşımı', level=2)
add_bullet(doc, 'Sözleşmesel alacak talepleri: 5 yıl (TBK Madde 146)')
add_bullet(doc, 'Haksız fiil tazminat talepleri: 2 yıl (zararın öğrenilmesinden), her hâlükârda 10 yıl (TBK Madde 72)')
add_bullet(doc, 'Kamu alacakları: 6183 Sayılı AATUHK uyarınca 5 yıl')
add_bullet(doc, 'KVKK ihlali şikâyeti: ihlalden itibaren 2 yıl (KVKK Madde 15)')

# ────────────────────────────
# BÖLÜM 12 — İSG
# ────────────────────────────
doc.add_page_break()
add_section_break_line(doc)
h = doc.add_heading('BÖLÜM 12 — İŞ SAĞLIĞI VE GÜVENLİĞİ YÜKÜMLÜLÜKLERİ', level=1)
h.runs[0].font.color.rgb = RGBColor(0x1F, 0x5C, 0x99)

doc.add_heading('12.1 Hukuki Dayanak', level=2)
add_body(doc, '6331 Sayılı İş Sağlığı ve Güvenliği Kanunu, İSG Risk Değerlendirmesi Yönetmeliği (RG: 29.12.2012/28512), Karayolu Taşıma Yönetmeliği Madde 75–82 (sürüş ve dinlenme süreleri) ile 5510 Sayılı Kanun\'un iş kazası hükümleri çerçevesinde düzenlenmiştir.')

doc.add_heading('12.2 İSG Eğitimleri', level=2)
isg_egitim = [
    ('Temel İSG Eğitimi (işe giriş)', 'En az 8 saat', 'İşe başlamadan önce', '6331 sk Madde 17'),
    ('Periyodik İSG Eğitimi', 'En az 4 saat/yıl', 'Her yıl', '6331 sk Madde 17'),
    ('Araç ve Ekipman Kullanım Eğitimi', 'En az 4 saat', 'Her yeni araç/ekipmanda', 'İSG Yönetmeliği'),
    ('İlk Yardım Eğitimi', 'En az 8 saat', 'İşe giriş + her 3 yılda bir', 'İlk Yardım Yönetmeliği'),
    ('Acil Tahliye ve Yangın Eğitimi', 'En az 4 saat/yıl', 'Her yıl', '6331 sk; Yangın Yönetmeliği'),
]
t = doc.add_table(rows=len(isg_egitim)+1, cols=4)
t.style = 'Table Grid'
add_header_row(t, ['Eğitim Türü', 'Süre', 'Sıklık', 'Dayanak'])
for i, row_data in enumerate(isg_egitim):
    add_table_row(t, i+1, row_data, alt_row=(i%2==0))
set_cell_borders(t)
doc.add_paragraph()

doc.add_heading('12.3 Sürüş Süreleri ve Dinlenme Zorunlulukları', level=2)
surus_data = [
    ('Günlük azami sürüş süresi', '9 saat', 'Haftada 2 gün 10 saate uzatılabilir'),
    ('Haftalık azami sürüş süresi', '56 saat', 'İki haftalık toplam 90 saati geçemez'),
    ('Kesintisiz azami sürüş süresi', '4,5 saat', 'Ardından en az 45 dk mola'),
    ('Günlük dinlenme süresi', 'En az 11 saat', '24 saatlik periyotta'),
    ('Haftalık dinlenme süresi', 'En az 45 saat', 'Her 6 günlük çalışma sonrası'),
    ('Kısa mola', 'En az 45 dakika (veya 15+30 dk bölünebilir)', '4,5 saatlik sürüş sonrası'),
]
t = doc.add_table(rows=len(surus_data)+1, cols=3)
t.style = 'Table Grid'
add_header_row(t, ['Parametre', 'Sınır', 'Açıklama'])
for i, row_data in enumerate(surus_data):
    add_table_row(t, i+1, row_data, alt_row=(i%2==0))
set_cell_borders(t)
doc.add_paragraph()
add_note(doc, 'Bu sınırlar akıllı takograf kayıtlarıyla belgelenir. Takograf kayıtları 365 gün saklanır ve talep halinde idareye sunulur.')

doc.add_heading('12.4 İş Kazası ve Meslek Hastalığı Bildirimi', level=2)
kaza_data = [
    ('SGK e-bildirge (iş kazası)', '3 iş günü içinde (ölümlü kazada derhal)', 'Sosyal Güvenlik Kurumu', '5510 sk Madde 13'),
    ('Çalışma ve Sosyal Güvenlik Bakanlığı İl Müdürlüğü', '48 saat içinde', 'İl Müdürlüğü (yazılı)', '6331 sk Madde 14'),
    ('İdare (ilgili birim)', '24 saat içinde', 'Hizmet alımı yürüten birim', 'Sözleşme hükmü'),
    ('Olay raporu (detaylı)', '72 saat içinde', 'İdare + SGK', '6331 sk; 5510 sk'),
]
t = doc.add_table(rows=len(kaza_data)+1, cols=4)
t.style = 'Table Grid'
add_header_row(t, ['Bildirim', 'Süre', 'Makam', 'Dayanak'])
for i, row_data in enumerate(kaza_data):
    add_table_row(t, i+1, row_data, alt_row=(i%2==0))
set_cell_borders(t)
doc.add_paragraph()

doc.add_heading('12.5 Araç Güvenlik Ekipmanları (Zorunlu)', level=2)
ekipman_data = [
    ('Yangın söndürücü', 'TS EN 3 uyumlu, en az 2 kg kuru toz', '6 ayda bir'),
    ('İlk yardım kiti', 'TS 11827 standardı', 'Her yıl veya kullanımdan sonra'),
    ('İkaz üçgeni', 'ECE R27 onaylı', 'Her araçta en az 1 adet'),
    ('Emniyet kemeri', 'ECE R16 onaylı (her koltukta)', 'Yıllık muayenede kontrol'),
    ('Araç yardım seti', 'Triger, tow halat, lastik tamiri', 'Mevcut olması yeterli'),
]
t = doc.add_table(rows=len(ekipman_data)+1, cols=3)
t.style = 'Table Grid'
add_header_row(t, ['Ekipman', 'Standart', 'Periyodik Kontrol'])
for i, row_data in enumerate(ekipman_data):
    add_table_row(t, i+1, row_data, alt_row=(i%2==0))
set_cell_borders(t)
doc.add_paragraph()

doc.add_heading('12.6 Denetim ve Raporlama', level=2)
add_bullet(doc, 'Yüklenici, aylık İSG faaliyet raporu hazırlar. Raporda dönemdeki kaza ve ramak kala olayları, gerçekleştirilen eğitimler, araç kontrol formlarının özeti ve alınan düzeltici önlemler yer alır.')
add_bullet(doc, 'Raporlar, her ayın son iş gününe kadar idareye elektronik ortamda iletilir.')
add_bullet(doc, 'İdare, önceden 24 saat bildirimle anlık denetim hakkını saklı tutar. Denetimde tespit edilen İSG eksiklikleri, 48 saat içinde giderilir.')

# ────────────────────────────
# REVIZYON GEÇMİŞİ
# ────────────────────────────
doc.add_page_break()
add_section_break_line(doc)
h = doc.add_heading('REVIZYON GEÇMİŞİ', level=1)
h.runs[0].font.color.rgb = RGBColor(0x1F, 0x5C, 0x99)

revizyon_data = [
    ('1.0', '24.04.2026', 'Teknik Şartname Uzmanı + Hukuk Müşaviri', 'İlk tam taslak — tüm bölümler tamamlandı'),
    ('1.1', '24.04.2026', 'Teknik Şartname Uzmanı', 'SRC-2→SRC-3 düzeltmesi; Bölüm 0 Tanımlar eklendi; KPI-11 eklendi; emisyon ve sigorta detaylandırıldı'),
    ('1.1', '24.04.2026', 'Hukuk Müşaviri', 'TÜFE güncelleme mekanizmaları (Bölüm 8.2, Bölüm 10.2) eklendi; mevzuat tablosu genişletildi'),
    ('1.1 — Final', '24.04.2026', 'Kalite Kontrolcü', 'Çapraz tutarlılık kontrolü tamamlandı; Word belgesi oluşturuldu'),
]
t = doc.add_table(rows=len(revizyon_data)+1, cols=4)
t.style = 'Table Grid'
add_header_row(t, ['Versiyon', 'Tarih', 'Güncelleyen', 'Değişiklik Özeti'])
for i, row_data in enumerate(revizyon_data):
    add_table_row(t, i+1, row_data, alt_row=(i%2==0))
set_cell_borders(t)
doc.add_paragraph()

# ────────────────────────────
# EKLER
# ────────────────────────────
add_section_break_line(doc)
h = doc.add_heading('EKLER', level=1)
h.runs[0].font.color.rgb = RGBColor(0x1F, 0x5C, 0x99)

ekler = [
    ('EK-1', 'Güzergah Haritaları ve Durak Listesi', 'Sözleşme imzalanmadan önce yüklenici tarafından sunulur'),
    ('EK-2', 'Araç Kabul Kriterleri Formu', 'Her araç için ayrı ayrı doldurulur'),
    ('EK-3', 'Araç Günlük Kontrol Formu', 'Her sefer öncesi ve sonrası sürücü tarafından imzalanır'),
    ('EK-4', 'Sürücü Bildirim Formu', 'Sürücü değişikliği, kaza, sağlık raporu bildirimleri için'),
    ('EK-5', 'Veri İşleme Sözleşmesi (DPA)', 'KVKK kapsamında; hizmet başlamadan önce imzalanır'),
    ('EK-6', 'Yolcu Memnuniyet Anketi', 'Üç aylık dönemlerde idare tarafından uygulanır'),
    ('EK-7', 'İSG Risk Değerlendirme Formu', 'Hizmet başlamadan önce yüklenici tarafından sunulur'),
    ('EK-8', 'Sefer Raporu Şablonu', 'Günlük ve aylık raporlar için standart format'),
]
t = doc.add_table(rows=len(ekler)+1, cols=3)
t.style = 'Table Grid'
add_header_row(t, ['Ek No', 'Başlık', 'Açıklama'])
for i, row_data in enumerate(ekler):
    add_table_row(t, i+1, row_data, alt_row=(i%2==0))
set_cell_borders(t)
doc.add_paragraph()

# ────────────────────────────
# ONAY BÖLÜMÜ
# ────────────────────────────
add_section_break_line(doc)
h = doc.add_heading('ONAY', level=1)
h.runs[0].font.color.rgb = RGBColor(0x1F, 0x5C, 0x99)

add_body(doc, 'Bu teknik şartname, aşağıdaki birimler tarafından incelenmiş ve onaylanmıştır.')
doc.add_paragraph()

onay_data = [
    ('Teknik Şartname Uzmanı', '24.04.2026', 'Bölüm 0–6 tamamlandı'),
    ('Hukuk Müşaviri', '24.04.2026', 'Bölüm 7–12 tamamlandı'),
    ('Kalite Kontrolcü', '24.04.2026', 'Tüm bölümler onaylı — belge yayına hazır'),
]
t = doc.add_table(rows=len(onay_data)+1, cols=3)
t.style = 'Table Grid'
add_header_row(t, ['Birim / Rol', 'Tarih', 'Durum'])
for i, row_data in enumerate(onay_data):
    add_table_row(t, i+1, row_data, alt_row=(i%2==0))
set_cell_borders(t)
doc.add_paragraph()

# Son not
add_note(doc, 'Bu belge Ajan Takımı (teknik-sartname-uzmani + hukuk-musaviri + kalite-kontrolcu) tarafından 24.04.2026 tarihinde üretilmiştir. Tüm hukuki atıflar Türk Hukuku kapsamındadır ve 2026 yılı itibarıyla yürürlükteki mevzuata dayanmaktadır.')

# ────────────────────────────
# KAYDET
# ────────────────────────────
output_path = r'C:\Users\furka\Desktop\sunum yz düzeltilecekler\agent teams eğitimi\personel-servisi-sartname.docx'
doc.save(output_path)
print(f"Belge basariyla olusturuldu: {output_path}")
